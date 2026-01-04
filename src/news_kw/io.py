"""Input/output operations for loading TXT and PDF news articles."""

import re
import warnings
import subprocess
import shutil
import tempfile
import os
from pathlib import Path
from typing import Optional, List
from concurrent.futures import ProcessPoolExecutor, as_completed
import pandas as pd
from tqdm import tqdm

try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    warnings.warn("pdfplumber not available, PDF support disabled")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    warnings.warn("python-docx not available, DOCX support disabled")


def extract_text_from_html(html_path: Path) -> Optional[str]:
    """Extract text from HTML file, removing JavaScript and CSS.
    
    Args:
        html_path: Path to HTML file
        
    Returns:
        Extracted text or None if error
    """
    try:
        with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Remove script tags and their content
        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove style tags and their content
        content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove HTML comments
        content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)
        
        # Remove HTML tags but keep text content
        content = re.sub(r'<[^>]+>', ' ', content)
        
        # Decode HTML entities (basic ones)
        content = content.replace('&nbsp;', ' ')
        content = content.replace('&amp;', '&')
        content = content.replace('&lt;', '<')
        content = content.replace('&gt;', '>')
        content = content.replace('&quot;', '"')
        content = content.replace('&#39;', "'")
        
        # Remove extra whitespace
        content = re.sub(r'\s+', ' ', content)
        
        return content.strip()
    except Exception as e:
        warnings.warn(f"Error extracting text from HTML {html_path}: {e}")
        return None


def parse_date_from_text(text: str, file_path: Optional[Path] = None, preferred_date: Optional[str] = None) -> Optional[str]:
    """Extract date from text content (various formats).
    
    Args:
        text: Full text content
        file_path: Optional file path (for year inference)
        preferred_date: Optional preferred date in YYYY-MM-DD format to match
        
    Returns:
        Date string in YYYY-MM-DD format or None
    """
    if not text:
        return None
    
    # If preferred_date is provided, try to find matching date first
    if preferred_date:
        # Try to find the preferred date in various formats
        preferred_year, preferred_month, preferred_day = preferred_date.split('-')
        
        # Check all date patterns and return if matches preferred_date
        patterns_to_check = [
            # YYYY-MM-DD
            (rf'{preferred_year}-{preferred_month}-{preferred_day}', lambda m: f"{m.group(1)}-{m.group(2)}-{m.group(3)}"),
            # Month DD, YYYY (need to convert month number to name)
            (rf'([A-Za-z]+\.?)\s+{preferred_day},?\s+{preferred_year}', None),  # Will check month separately
            # M/D/YYYY or M/D/YY
            (rf'{int(preferred_month)}/{int(preferred_day)}/{preferred_year}', None),
            (rf'{int(preferred_month)}/{int(preferred_day)}/{preferred_year[2:]}', None),
        ]
        
        # Check for Month DD, YYYY format matching preferred date
        month_names = {
            'january': '01', 'jan': '01', 'jan.': '01',
            'february': '02', 'feb': '02', 'feb.': '02',
            'march': '03', 'mar': '03', 'mar.': '03',
            'april': '04', 'apr': '04', 'apr.': '04',
            'may': '05', 'may.': '05',
            'june': '06', 'jun': '06', 'jun.': '06',
            'july': '07', 'jul': '07', 'jul.': '07',
            'august': '08', 'aug': '08', 'aug.': '08',
            'september': '09', 'sep': '09', 'sep.': '09', 'sept': '09', 'sept.': '09',
            'october': '10', 'oct': '10', 'oct.': '10',
            'november': '11', 'nov': '11', 'nov.': '11',
            'december': '12', 'dec': '12', 'dec.': '12'
        }
        
        # Reverse lookup: find month name for preferred_month
        month_name_candidates = [name for name, num in month_names.items() if num == preferred_month]
        
        if month_name_candidates:
            for month_name in month_name_candidates:
                # Try various formats
                patterns = [
                    rf'{month_name}\s+{int(preferred_day)},?\s+{preferred_year}',
                    rf'{month_name}\.\s+{int(preferred_day)},?\s+{preferred_year}',
                ]
                for pattern in patterns:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        return preferred_date
        
        # Try YYYY-MM-DD format
        if preferred_date in text:
            return preferred_date
        
        # Try M/D/YYYY format
        md_pattern = rf'{int(preferred_month)}/{int(preferred_day)}/{preferred_year}'
        if re.search(md_pattern, text):
            return preferred_date
    
    # Month names mapping
    month_names = {
        'january': '01', 'jan': '01', 'jan.': '01',
        'february': '02', 'feb': '02', 'feb.': '02',
        'march': '03', 'mar': '03', 'mar.': '03',
        'april': '04', 'apr': '04', 'apr.': '04',
        'may': '05', 'may.': '05',
        'june': '06', 'jun': '06', 'jun.': '06',
        'july': '07', 'jul': '07', 'jul.': '07',
        'august': '08', 'aug': '08', 'aug.': '08',
        'september': '09', 'sep': '09', 'sep.': '09', 'sept': '09', 'sept.': '09',
        'october': '10', 'oct': '10', 'oct.': '10',
        'november': '11', 'nov': '11', 'nov.': '11',
        'december': '12', 'dec': '12', 'dec.': '12'
    }
    
    # 1. Try "Date: YYYY-MM-DD" format
    match = re.search(r'Date:\s*(\d{4}-\d{2}-\d{2})', text, re.IGNORECASE)
    if match:
        return match.group(1)
    
    # 2. Try "M/D/YY" or "M/D/YYYY" format (e.g., "1/14/20", "1/14/2020", "01/14/20")
    # This is a common US date format and should be checked early
    # This format is often the actual document date
    md_yy_match = re.search(r'(\d{1,2})/(\d{1,2})/(\d{2,4})', text)
    if md_yy_match:
        month, day, year = md_yy_match.groups()
        month_int = int(month)
        day_int = int(day)
        year_int = int(year)
        
        # Validate month and day
        if 1 <= month_int <= 12 and 1 <= day_int <= 31:
            # Handle 2-digit year: assume 2000s if year < 50, 1900s if year >= 50
            if year_int < 100:
                if year_int < 50:
                    year_int = 2000 + year_int
                else:
                    year_int = 1900 + year_int
            
            month_padded = str(month_int).zfill(2)
            day_padded = str(day_int).zfill(2)
            return f"{year_int}-{month_padded}-{day_padded}"
    
    # 3. Try "DD Month YYYY" format (e.g., "19 December 2018")
    # Look for pattern: 1-2 digits, month name, 4-digit year
    month_day_year = re.search(
        r'(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})',
        text,
        re.IGNORECASE
    )
    if month_day_year:
        day, month_str, year = month_day_year.groups()
        month_lower = month_str.lower()
        if month_lower in month_names:
            month = month_names[month_lower]
            day_padded = day.zfill(2)
            return f"{year}-{month}-{day_padded}"
    
    # 4. Try "Month DD, YYYY" or "Month DD YYYY" format (e.g., "December 19, 2018" or "Oct. 09, 2018")
    # Support both "Month." and "Month" formats
    month_day_year2 = re.search(
        r'([A-Za-z]+\.?)\s+(\d{1,2}),?\s+(\d{4})',
        text,
        re.IGNORECASE
    )
    if month_day_year2:
        month_str, day, year = month_day_year2.groups()
        month_lower = month_str.lower().rstrip('.')  # Remove trailing dot if present
        if month_lower in month_names:
            month = month_names[month_lower]
            day_padded = day.zfill(2)
            return f"{year}-{month}-{day_padded}"
    
    # 4. Try "Updated: Month DD, YYYY" or "Published: Month DD, YYYY" format
    # (e.g., "Updated: Oct. 09, 2018" or "Published: Oct. 09, 2018")
    updated_published = re.search(
        r'(?:Updated|Published):\s*([A-Za-z]+\.?)\s+(\d{1,2}),?\s+(\d{4})',
        text,
        re.IGNORECASE
    )
    if updated_published:
        month_str, day, year = updated_published.groups()
        month_lower = month_str.lower().rstrip('.')  # Remove trailing dot if present
        if month_lower in month_names:
            month = month_names[month_lower]
            day_padded = day.zfill(2)
            return f"{year}-{month}-{day_padded}"
    
    # 5. Try "Broadcast: Weekday, Month DD, YYYY" or "Broadcast: Weekday, Month DD" format
    # (e.g., "Broadcast: Tuesday, Aug. 17, 2021" or "Broadcast: Tuesday, Aug. 17")
    broadcast_with_year = re.search(
        r'Broadcast:\s*[A-Za-z]+,\s*([A-Za-z]+\.?)\s+(\d{1,2}),?\s+(\d{4})',
        text,
        re.IGNORECASE
    )
    if broadcast_with_year:
        month_str, day, year = broadcast_with_year.groups()
        month_lower = month_str.lower().rstrip('.')  # Remove trailing dot if present
        if month_lower in month_names:
            month = month_names[month_lower]
            day_padded = day.zfill(2)
            return f"{year}-{month}-{day_padded}"
    
    # 6. Try "Broadcast: Weekday, Month DD" format (year from filename will be used)
    # (e.g., "Broadcast: Tuesday, Aug. 17")
    # Note: This returns a special format that will be combined with year from filename
    broadcast_no_year = re.search(
        r'Broadcast:\s*[A-Za-z]+,\s*([A-Za-z]+\.?)\s+(\d{1,2})(?!\s*\d{4})',
        text,
        re.IGNORECASE
    )
    if broadcast_no_year:
        month_str, day = broadcast_no_year.groups()
        month_lower = month_str.lower().rstrip('.')  # Remove trailing dot if present
        if month_lower in month_names:
            month = month_names[month_lower]
            day_padded = day.zfill(2)
            # Return special format "PARTIAL-MM-DD" to be combined with year from filename
            # This will be handled in validate_date_parsing
            return f"PARTIAL-{month}-{day_padded}"
    
    # 7. Try "YYYY-MM-DD" format anywhere in text
    ymd_match = re.search(r'(\d{4})-(\d{2})-(\d{2})', text)
    if ymd_match:
        year, month, day = ymd_match.groups()
        # Validate
        if 1 <= int(month) <= 12 and 1 <= int(day) <= 31:
            return f"{year}-{month}-{day}"
    
    return None


def parse_date_from_path(file_path: Path) -> Optional[str]:
    """Extract date from file path (filename only, no folder structure).
    
    Only extracts date from filename. If filename has no date information, returns None.
    
    Priority: Filename prefix (YYYY-MM-DD_) > Other filename patterns
    
    Args:
        file_path: Path to the file
        
    Returns:
        Date string in YYYY-MM-DD format or None if no date found in filename
    """
    # 0. Check filename prefix for YYYY-MM-DD_ format (highest priority)
    # This is the standard format used in filtered_data
    prefix_match = re.match(r'^(\d{4})-(\d{2})-(\d{2})_', file_path.name)
    if prefix_match:
        year, month, day = prefix_match.groups()
        # Validate date
        if 1 <= int(month) <= 12 and 1 <= int(day) <= 31:
            return f"{year}-{month}-{day}"
    
    # 1. Check filename for YYYY-MM-DD, YYYY_MM_DD, or YYYY.MM.DD
    filename_match = re.search(r'(\d{4})[-_.](\d{2})[-_.](\d{2})', file_path.name)
    if filename_match:
        year, month, day = filename_match.groups()
        return f"{year}-{month}-{day}"
    
    # 1.5. Check filename for MM_DD_YYYY or MM-DD-YYYY format (e.g., "02_24_2022", "12_06_2022")
    # This pattern appears in some news article filenames
    mm_dd_yyyy_match = re.search(r'(\d{1,2})[-_](\d{1,2})[-_](\d{4})', file_path.name)
    if mm_dd_yyyy_match:
        num1, num2, year = mm_dd_yyyy_match.groups()
        num1_int = int(num1)
        num2_int = int(num2)
        # Validate: first number should be month (1-12), second should be day (1-31)
        if 1 <= num1_int <= 12 and 1 <= num2_int <= 31:
            month = num1.zfill(2)
            day = num2.zfill(2)
            return f"{year}-{month}-{day}"
        # Try reverse order: DD_MM_YYYY (if first number > 12, it might be day)
        elif 1 <= num2_int <= 12 and 1 <= num1_int <= 31:
            month = num2.zfill(2)
            day = num1.zfill(2)
            return f"{year}-{month}-{day}"
    
    # 1.6. Check filename for MM_DD_YY or MM-DD-YY format (2-digit year, e.g., "04_17_23", "12_06_22")
    # This pattern appears in some news article filenames with 2-digit year
    mm_dd_yy_match = re.search(r'(\d{1,2})[-_](\d{1,2})[-_](\d{2})(?!\d)', file_path.name)
    if mm_dd_yy_match:
        num1, num2, year_2digit = mm_dd_yy_match.groups()
        num1_int = int(num1)
        num2_int = int(num2)
        year_int = int(year_2digit)
        # Handle 2-digit year: assume 2000s if year < 50, 1900s if year >= 50
        if year_int < 50:
            year_int = 2000 + year_int
        else:
            year_int = 1900 + year_int
        year = str(year_int)
        # Validate: first number should be month (1-12), second should be day (1-31)
        if 1 <= num1_int <= 12 and 1 <= num2_int <= 31:
            month = num1.zfill(2)
            day = num2.zfill(2)
            return f"{year}-{month}-{day}"
        # Try reverse order: DD_MM_YY (if first number > 12, it might be day)
        elif 1 <= num2_int <= 12 and 1 <= num1_int <= 31:
            month = num2.zfill(2)
            day = num1.zfill(2)
            return f"{year}-{month}-{day}"
    
    # 2. Check filename for MMM. DD, YYYY format (e.g., "Nov. 07, 2018")
    month_names = {
        'jan': '01', 'january': '01', 'jan.': '01',
        'feb': '02', 'february': '02', 'feb.': '02',
        'mar': '03', 'march': '03', 'mar.': '03',
        'apr': '04', 'april': '04', 'apr.': '04',
        'may': '05', 'may.': '05',
        'jun': '06', 'june': '06', 'jun.': '06',
        'jul': '07', 'july': '07', 'jul.': '07',
        'aug': '08', 'august': '08', 'aug.': '08',
        'sep': '09', 'september': '09', 'sep.': '09', 'sept.': '09',
        'oct': '10', 'october': '10', 'oct.': '10',
        'nov': '11', 'november': '11', 'nov.': '11',
        'dec': '12', 'december': '12', 'dec.': '12'
    }
    
    # Pattern: MMM. DD, YYYY or MMM DD, YYYY or MMM DD_YYYY (e.g., "Nov. 07, 2018", "July 17_2020")
    month_day_year = re.search(
        r'([A-Za-z]+\.?)\s+(\d{1,2})[,_\s]+(\d{4})', 
        file_path.name, 
        re.IGNORECASE
    )
    if month_day_year:
        month_str, day, year = month_day_year.groups()
        month_lower = month_str.lower().rstrip('.')  # Remove trailing dot if present
        if month_lower in month_names:
            month = month_names[month_lower]
            day_padded = day.zfill(2)
            return f"{year}-{month}-{day_padded}"
    
    # 3. Check for MMM_YYYY or MMM-YYYY format in filename (e.g., "Feb_2022", "Feb-2022")
    month_year_pattern = re.search(
        r'([A-Za-z]+)[-_](\d{4})',
        file_path.name,
        re.IGNORECASE
    )
    if month_year_pattern:
        month_str, year = month_year_pattern.groups()
        month_lower = month_str.lower().rstrip('.')  # Remove trailing dot if present
        if month_lower in month_names:
            month = month_names[month_lower]
            # Use first day of month
            return f"{year}-{month}-01"
    
    # 3.5. Check for MM_YYYY or MM-YYYY format in filename (numeric month only, e.g., "09_2024", "06_2025")
    # This pattern appears in some filenames with numeric month
    mm_yyyy_match = re.search(r'(\d{1,2})[-_](\d{4})(?![-_]\d)', file_path.name)
    if mm_yyyy_match:
        month_str, year = mm_yyyy_match.groups()
        month_int = int(month_str)
        # Validate month
        if 1 <= month_int <= 12:
            month = month_str.zfill(2)
            # Use first day of month
            return f"{year}-{month}-01"
    
    # 4. Check for YYYY-MM, YYYY_MM, or YYYY.MM in filename
    year_month = re.search(r'(\d{4})[-_.](\d{2})(?![-_.]\d{2})', file_path.name)
    if year_month:
        year, month = year_month.groups()
        return f"{year}-{month}-01"
    
    # No date information found in filename - return None
    return None


def parse_title_from_text(text: str) -> Optional[str]:
    """Extract title from text header (Title: ... format).
    
    Args:
        text: Full text content
        
    Returns:
        Title string or None
    """
    match = re.search(r'Title:\s*(.+?)(?:\n|$)', text, re.IGNORECASE | re.MULTILINE)
    if match:
        return match.group(1).strip()
    return None


def parse_source_from_text(text: str) -> Optional[str]:
    """Extract source from text header (Source: ... format).
    
    Args:
        text: Full text content
        
    Returns:
        Source string or None
    """
    match = re.search(r'Source:\s*(.+?)(?:\n|$)', text, re.IGNORECASE | re.MULTILINE)
    if match:
        return match.group(1).strip()
    return None


def extract_text_from_pdf(pdf_path: Path) -> Optional[str]:
    """Extract text from PDF file.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Extracted text or None if error
    """
    if not PDF_SUPPORT:
        return None
    
    try:
        text_parts = []
        # Suppress pdfplumber warnings about invalid PDF metadata (fonts, colors, etc.)
        # These warnings don't affect text extraction
        with warnings.catch_warnings():
            warnings.filterwarnings('ignore', category=UserWarning, module='pdfplumber')
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        return '\n'.join(text_parts)
    except Exception as e:
        warnings.warn(f"Error extracting text from PDF {pdf_path}: {e}")
        return None


def extract_text_from_docx(docx_path: Path) -> Optional[str]:
    """Extract text from DOCX file.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Extracted text or None if error
    """
    if not DOCX_SUPPORT:
        return None
    
    try:
        doc = Document(docx_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n'.join(text_parts) if text_parts else None
    except Exception as e:
        warnings.warn(f"Error extracting text from DOCX {docx_path}: {e}")
        return None


def convert_docx_to_pdf(docx_path: Path, output_pdf_path: Path) -> bool:
    """Convert DOCX file to PDF using LibreOffice or docx2pdf.
    
    Args:
        docx_path: Path to DOCX file
        output_pdf_path: Path to save converted PDF file
        
    Returns:
        True if conversion successful, False otherwise
    """
    # Try LibreOffice first (cross-platform, free)
    libreoffice_path = shutil.which('soffice') or shutil.which('libreoffice')
    if libreoffice_path:
        try:
            # LibreOffice command: soffice --headless --convert-to pdf --outdir <dir> <file>
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run(
                    [libreoffice_path, '--headless', '--convert-to', 'pdf', 
                     '--outdir', tmpdir, str(docx_path)],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0:
                    # Find the converted PDF (LibreOffice uses original filename)
                    converted_pdf = Path(tmpdir) / f"{docx_path.stem}.pdf"
                    if converted_pdf.exists():
                        shutil.copy2(converted_pdf, output_pdf_path)
                        return True
        except Exception as e:
            warnings.warn(f"LibreOffice conversion failed for {docx_path}: {e}")
    
    # Try docx2pdf (Windows-friendly, requires MS Word)
    try:
        import docx2pdf
        docx2pdf.convert(str(docx_path), str(output_pdf_path))
        if output_pdf_path.exists():
            return True
    except ImportError:
        pass
    except Exception as e:
        warnings.warn(f"docx2pdf conversion failed for {docx_path}: {e}")
    
    return False


def extract_text_from_docx_with_fallback(docx_path: Path) -> Optional[str]:
    """Extract text from DOCX file, fallback to PDF conversion if needed.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Extracted text or None if error
    """
    # First try direct DOCX reading
    text = extract_text_from_docx(docx_path)
    if text and text.strip():
        return text
    
    # If direct reading failed or returned empty, try PDF conversion
    if not PDF_SUPPORT:
        warnings.warn(f"Could not extract text from DOCX {docx_path} and PDF support is disabled")
        return None
    
    warnings.warn(f"Could not extract text directly from DOCX {docx_path}, attempting PDF conversion...")
    
    # Convert to PDF in temporary location
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = Path(tmpdir) / f"{docx_path.stem}.pdf"
        if convert_docx_to_pdf(docx_path, pdf_path):
            # Extract text from converted PDF
            return extract_text_from_pdf(pdf_path)
        else:
            warnings.warn(f"Could not convert DOCX {docx_path} to PDF")
            return None


def _process_single_file(file_path: Path) -> Optional[dict]:
    """Process a single file and return document dict or None (for parallel processing).
    
    Args:
        file_path: Path to the file to process
        
    Returns:
        Document dict or None if processing failed
    """
    try:
        # Extract text based on file type
        file_ext = file_path.suffix.lower()
        if file_ext == '.pdf':
            content = extract_text_from_pdf(file_path)
            if not content:
                return None
        elif file_ext == '.docx':
            content = extract_text_from_docx_with_fallback(file_path)
            if not content:
                return None
        elif file_ext in ['.html', '.htm']:
            content = extract_text_from_html(file_path)
            if not content:
                return None
        else:
            # Assume TXT file
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        
        # Parse date from filename prefix only (YYYY-MM-DD_ format)
        # This is the standard format used in filtered_data directory
        date = parse_date_from_path(file_path)
        if not date:
            # If no prefix date found, skip this file
            warnings.warn(f"No date prefix (YYYY-MM-DD_) found in filename for {file_path}, skipping")
            return None
        
        # Parse title
        title = parse_title_from_text(content)
        if not title:
            title = file_path.stem
        
        # Parse source
        source = parse_source_from_text(content)
        if not source:
            source = "unknown"
        
        # Extract body
        text = extract_body_text(content)
        
        if not text.strip():
            return None
        
        doc_id = f"{file_path.stem}_{file_path.parent.name}"
        
        return {
            'doc_id': doc_id,
            'date': date,
            'title': title,
            'text': text,
            'source': source
        }
    except Exception as e:
        warnings.warn(f"Error processing {file_path}: {e}")
        return None


def extract_body_text(text: str) -> str:
    """Extract body text by removing header lines.
    
    Args:
        text: Full text content
        
    Returns:
        Body text without headers
    """
    lines = text.split('\n')
    body_lines = []
    skip_headers = False
    
    for line in lines:
        # Skip header lines
        if re.match(r'^(Title|Date|Source):', line, re.IGNORECASE):
            skip_headers = True
            continue
        
        # After headers, start collecting body
        if skip_headers and line.strip():
            body_lines.append(line)
        elif not skip_headers:
            body_lines.append(line)
    
    return '\n'.join(body_lines).strip()


def check_files_without_prefix_date(all_files: List[Path]) -> List[Path]:
    """Check for files without date prefix in filename.
    
    Args:
        all_files: List of file paths to check (already filtered by source folders)
        
    Returns:
        List of file paths without date prefix
    """
    files_without_prefix = []
    
    # Check each file for date prefix
    for file_path in all_files:
        # Check if filename starts with YYYY-MM-DD_ format
        if not re.match(r'^\d{4}-\d{2}-\d{2}_', file_path.name):
            files_without_prefix.append(file_path)
    
    return files_without_prefix


def load_txt_articles(input_dir: Path, output_dir: Path, source_folders: Optional[List[str]] = None) -> pd.DataFrame:
    """Load all TXT, PDF, and DOCX articles from directory recursively.
    
    Args:
        input_dir: Directory containing TXT, PDF, and DOCX files
        output_dir: Directory to save processed documents
        source_folders: List of folder names to read from (e.g., ['meeting', 'news', 'raddit']).
                       If None, reads from all subdirectories.
        
    Returns:
        DataFrame with columns: doc_id, date, title, text, source
    """
    # Get all supported files recursively (rglob searches all subdirectories)
    # This will find files in: folder/file.txt, folder/subfolder/file.txt, folder/sub1/sub2/file.txt, etc.
    txt_files = list(input_dir.rglob('*.txt'))
    pdf_files = list(input_dir.rglob('*.pdf')) if PDF_SUPPORT else []
    docx_files = list(input_dir.rglob('*.docx')) + list(input_dir.rglob('*.DOCX'))
    
    all_files = txt_files + pdf_files + docx_files
    
    # Filter by source folders if specified
    # Note: This checks only the first folder in the path, so files in subdirectories
    # (e.g., raddit/2018/file.pdf, raddit/2019/subfolder/file.pdf) are all included
    # as long as the top-level folder matches one of the source_folders
    if source_folders:
        filtered_files = []
        for file_path in all_files:
            # Check if file is in any of the specified source folders
            # Get relative path from input_dir
            try:
                rel_path = file_path.relative_to(input_dir)
                # Get the first part of the path (top-level folder name)
                # This allows matching files in any depth of subdirectories
                # Example: raddit/2018/file.pdf -> first_folder = 'raddit' ✓
                #          raddit/2019/subfolder/file.pdf -> first_folder = 'raddit' ✓
                first_folder = rel_path.parts[0] if rel_path.parts else None
                if first_folder in source_folders:
                    filtered_files.append(file_path)
            except ValueError:
                # File is not under input_dir, skip it
                continue
        all_files = filtered_files
    
    if not all_files:
        folder_info = f" in folders {source_folders}" if source_folders else ""
        raise ValueError(f"No TXT, PDF, or DOCX files found in {input_dir}{folder_info}")
    
    # Check for files without date prefix and create report
    # Reuse the already-scanned file list to avoid duplicate scanning
    files_without_prefix = check_files_without_prefix_date(all_files)
    if files_without_prefix:
        # Create report file
        report_path = output_dir.parent / 'files_without_prefix_date.txt'
        report_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("# Files without date prefix (YYYY-MM-DD_) in filename\n")
            f.write(f"# Total files without prefix: {len(files_without_prefix)}\n")
            f.write(f"# Total files checked: {len(all_files)}\n")
            f.write("#\n")
            f.write("# These files do not have the expected date prefix format.\n")
            f.write("# Expected format: YYYY-MM-DD_filename.pdf\n")
            f.write("#\n\n")
            
            for file_path in files_without_prefix:
                # Get relative path from input_dir
                try:
                    rel_path = file_path.relative_to(input_dir)
                    f.write(f"{rel_path}\n")
                except ValueError:
                    f.write(f"{file_path}\n")
        
        warnings.warn(
            f"경고: {len(files_without_prefix)}개 파일에 날짜 prefix (YYYY-MM-DD_)가 없습니다. "
            f"리포트 파일: {report_path}"
        )
    
    # Process files in parallel
    cpu_count = os.cpu_count() or 1
    max_workers = max(1, int(cpu_count * 0.7))
    num_files = len(all_files)
    workers = min(max_workers, num_files)
    
    documents = []
    
    if num_files > 10 and workers > 1:
        # Parallel processing for large file sets
        processed_files = set()
        failed_files = []
        skipped_files = []
        
        with ProcessPoolExecutor(max_workers=workers) as executor:
            futures = {executor.submit(_process_single_file, file_path): file_path for file_path in all_files}
            
            # Process all futures and track results
            for future in tqdm(as_completed(futures), total=len(futures), desc="Loading files"):
                file_path = futures[future]
                try:
                    result = future.result()
                    if result:
                        documents.append(result)
                        processed_files.add(file_path)
                    else:
                        # File was processed but returned None (skipped due to validation issues)
                        skipped_files.append(file_path)
                except Exception as e:
                    # File processing failed with exception
                    failed_files.append((file_path, str(e)))
                    warnings.warn(f"Failed to process file {file_path}: {e}")
        
        # Verify all files were processed
        total_processed = len(processed_files) + len(skipped_files) + len(failed_files)
        if total_processed != num_files:
            missing_count = num_files - total_processed
            warnings.warn(
                f"파일 처리 누락 경고: {missing_count}개 파일이 처리되지 않았습니다. "
                f"(전체: {num_files}, 처리됨: {len(processed_files)}, 스킵됨: {len(skipped_files)}, 실패: {len(failed_files)})"
            )
        
        # Log detailed statistics
        if failed_files or skipped_files:
            warnings.warn(
                f"파일 처리 요약: "
                f"성공 {len(processed_files)}개, "
                f"스킵 {len(skipped_files)}개, "
                f"실패 {len(failed_files)}개"
            )
            if failed_files:
                warnings.warn(f"실패한 파일 목록 (최대 10개):")
                for file_path, error in failed_files[:10]:
                    warnings.warn(f"  - {file_path}: {error}")
                if len(failed_files) > 10:
                    warnings.warn(f"  ... 외 {len(failed_files) - 10}개 파일 실패")
    else:
        # Sequential processing for small file sets
        for file_path in tqdm(all_files, desc="Loading files"):
            try:
                # Extract text based on file type
                file_ext = file_path.suffix.lower()
                if file_ext == '.pdf':
                    content = extract_text_from_pdf(file_path)
                    if not content:
                        warnings.warn(f"Could not extract text from PDF {file_path}, skipping")
                        continue
                elif file_ext == '.docx':
                    content = extract_text_from_docx_with_fallback(file_path)
                    if not content:
                        warnings.warn(f"Could not extract text from DOCX {file_path}, skipping")
                        continue
                else:
                    # Assume TXT file
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                
                # Parse date from filename prefix only (YYYY-MM-DD_ format)
                # This is the standard format used in filtered_data directory
                date = parse_date_from_path(file_path)
                if not date:
                    warnings.warn(f"No date prefix (YYYY-MM-DD_) found in filename for {file_path}, skipping")
                    continue
                
                # Parse title
                title = parse_title_from_text(content)
                if not title:
                    title = file_path.stem
                
                # Parse source
                source = parse_source_from_text(content)
                if not source:
                    source = "unknown"
                
                # Extract body
                text = extract_body_text(content)
                
                if not text.strip():
                    warnings.warn(f"Empty text for {file_path}, skipping")
                    continue
                
                doc_id = f"{file_path.stem}_{file_path.parent.name}"
                
                documents.append({
                    'doc_id': doc_id,
                    'date': date,
                    'title': title,
                    'text': text,
                    'source': source
                })
                
            except Exception as e:
                warnings.warn(f"Error processing {file_path}: {e}")
                continue
    
    if not documents:
        raise ValueError("No valid documents loaded")
    
    df = pd.DataFrame(documents)
    
    # Convert dates with error handling for invalid dates
    try:
        df['date'] = pd.to_datetime(df['date'], errors='coerce', format='mixed')
    except Exception as e:
        # Fallback: try with errors='coerce' to handle invalid dates
        warnings.warn(f"Error parsing dates, attempting with errors='coerce': {e}")
        df['date'] = pd.to_datetime(df['date'], errors='coerce')
    
    # Remove rows with invalid dates
    invalid_dates = df['date'].isna()
    if invalid_dates.any():
        invalid_count = invalid_dates.sum()
        warnings.warn(f"Removing {invalid_count} documents with invalid dates")
        df = df[~invalid_dates].copy()
    
    if len(df) == 0:
        raise ValueError("No valid documents with valid dates loaded")
    
    df = df.sort_values('date').reset_index(drop=True)
    
    # Save to parquet
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / 'documents.parquet'
    df.to_parquet(output_path, index=False)
    
    return df

